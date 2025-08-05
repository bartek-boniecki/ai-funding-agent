# main.py

from fastapi import FastAPI, Request
from fastapi.responses import FileResponse, JSONResponse
import uvicorn
import logging
import os
from dotenv import load_dotenv

# LangChain imports
from langchain.utilities import SerpAPIWrapper
from langchain.chat_models import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain

# Document generation
from docx import Document

# Email sending via SendGrid
from sendgrid import SendGridAPIClient
from sendgrid.helpers.mail import Mail, Attachment, FileContent, FileName, FileType, Disposition
import base64

# 1️⃣ Load environment variables
load_dotenv()

# Validate SendGrid configuration
SENDGRID_API_KEY = os.getenv("SENDGRID_API_KEY")
SENDGRID_SENDER = os.getenv("SENDGRID_SENDER")
if not SENDGRID_API_KEY or not SENDGRID_SENDER:
    raise RuntimeError("SENDGRID_API_KEY and SENDGRID_SENDER must be set in .env")

# 2️⃣ Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 3️⃣ Initialize FastAPI
app = FastAPI()

# 4️⃣ Initialize SerpAPI and a chat-model LLM
serp = SerpAPIWrapper(serpapi_api_key=os.getenv("SERPAPI_API_KEY"))
llm = ChatOpenAI(
    model_name="gpt-4",
    temperature=0,
    openai_api_key=os.getenv("OPENAI_API_KEY"),
    max_tokens=2048
)

# 5️⃣ Query-refinement chains
template_a = PromptTemplate(
    input_variables=["problem_text"],
    template=(
        "You are a search query generator. Given a user-provided problem description, "
        "produce a concise search query to find publicly available statistics illustrating "
        "the scale and urgency of the problem.\n\n"
        "Problem:\n{problem_text}\n\nSearch Query:"
    )
)
chain_a = LLMChain(llm=llm, prompt=template_a)

template_b = PromptTemplate(
    input_variables=["solution_text"],
    template=(
        "You are a search query generator. Given a product/solution description, "
        "produce a concise search query to find the market addressed, market size, "
        "and key trends.\n\n"
        "Product:\n{solution_text}\n\nSearch Query:"
    )
)
chain_b = LLMChain(llm=llm, prompt=template_b)

template_c = PromptTemplate(
    input_variables=["features_text"],
    template=(
        "You are a search query generator. Given unique features of a product, "
        "produce a concise search query to find companies offering similar features "
        "and assess the degree of novelty.\n\n"
        "Features:\n{features_text}\n\nSearch Query:"
    )
)
chain_c = LLMChain(llm=llm, prompt=template_c)

template_d = PromptTemplate(
    input_variables=["features_text"],
    template=(
        "You are a search query generator. Given unique features of a product, "
        "produce a concise search query to find competitor revenue streams.\n\n"
        "Features:\n{features_text}\n\nSearch Query:"
    )
)
chain_d = LLMChain(llm=llm, prompt=template_d)

# 6️⃣ Analysis chains
analysis_template_a = PromptTemplate(
    input_variables=["problem_text","search_results"],
    template=(
        "You are an expert grant writer. Draft a 1500–2000 character analysis titled 'The problem/market opportunity'.\n\n"
        "Problem description:\n{problem_text}\n\n"
        "Search snippets:\n{search_results}\n\n"
        "Include concrete statistics illustrating scale and urgency. Do NOT truncate mid-sentence."
    )
)
analysis_chain_a = LLMChain(llm=llm, prompt=analysis_template_a)

analysis_template_b = PromptTemplate(
    input_variables=["solution_text","search_results","trl","features_text"],
    template=(
        "Draft a 1500–2000 character analysis titled 'The innovation: Solution/Product or Services (USP)'.\n\n"
        "1. Why the product is better than existing solutions (use: {search_results}).\n"
        "2. Current TRL: {trl} (validation/certification details if any).\n"
        "3. Why now is the right time to bring this product to market.\n\n"
        "Do NOT truncate mid-sentence."
    )
)
analysis_chain_b = LLMChain(llm=llm, prompt=analysis_template_b)

analysis_template_c = PromptTemplate(
    input_variables=["solution_text","search_results","revenue_model","features_text"],
    template=(
        "Draft a 1500–2000 character analysis titled 'Market and Competition analysis'.\n\n"
        "1. Market size and key trends (use: {search_results}).\n"
        "2. Product potential to transform or create a market.\n"
        "3. Business model & revenue streams: {revenue_model}.\n"
        "4. Why unique features: {features_text} drive adoption.\n"
        "5. Advantages/disadvantages & likelihood of success.\n\n"
        "Do NOT truncate mid-sentence."
    )
)
analysis_chain_c = LLMChain(llm=llm, prompt=analysis_template_c)

analysis_template_d = PromptTemplate(
    input_variables=["solution_text","search_results"],
    template=(
        "Draft a 1500–2000 character analysis titled 'Broad impacts'.\n\n"
        "Discuss potential societal, environmental, or climate impacts and estimate job creation. Use: {search_results}.\n\n"
        "Do NOT truncate mid-sentence."
    )
)
analysis_chain_d = LLMChain(llm=llm, prompt=analysis_template_d)

analysis_template_e = PromptTemplate(
    input_variables=["solution_text","revenue_model","trl"],
    template=(
        "Draft a 1500–2000 character analysis titled 'Funding rationale and MVP'.\n\n"
        "Explain why raising funding is nearly impossible at TRL {trl}, why building an MVP is crucial, "
        "and include any funding history if provided. Use revenue model: {revenue_model}.\n\n"
        "Do NOT truncate mid-sentence."
    )
)
analysis_chain_e = LLMChain(llm=llm, prompt=analysis_template_e)

@app.post("/webhook")
async def receive_form(request: Request):
    payload = await request.json()
    questions = payload.get("submission", {}).get("questions", [])

    # 7️⃣ Extract inputs and email
    try:
        solution = questions[0].get("value", "")
        problem = questions[1].get("value", "")
        unique_feature = questions[2].get("value", "")
        current_trl = questions[3].get("value", "")
        revenue_model = questions[4].get("value", "")
        user_email = questions[5].get("value", "").strip()
    except Exception:
        return JSONResponse(status_code=400, content={"detail": "Form fields missing or misordered."})

    if not user_email:
        # Instruct user to include email if blank
        return JSONResponse(status_code=400, content={"detail": "Email address is required in the form."})

    logger.info("Inputs received; emailing to: %s", user_email)

    # 8️⃣ Generate search snippets
    query_a = chain_a.invoke({"problem_text": problem})
    snippets_a = "\n".join(serp.run(query_a)[:5])
    query_b = chain_b.invoke({"solution_text": solution})
    snippets_b = "\n".join(serp.run(query_b)[:5])
    query_c = chain_c.invoke({"features_text": unique_feature})
    snippets_c = "\n".join(serp.run(query_c)[:5])
    query_d = chain_d.invoke({"features_text": unique_feature})
    snippets_d = "\n".join(serp.run(query_d)[:5])

    # 9️⃣ Draft analyses
    analysis_a = analysis_chain_a.invoke({"problem_text": problem, "search_results": snippets_a})
    analysis_b = analysis_chain_b.invoke({"solution_text": solution, "search_results": snippets_b, "trl": current_trl, "features_text": unique_feature})
    analysis_c = analysis_chain_c.invoke({"solution_text": solution, "search_results": snippets_b, "revenue_model": revenue_model, "features_text": unique_feature})
    analysis_d = analysis_chain_d.invoke({"solution_text": solution, "search_results": snippets_d})
    analysis_e = analysis_chain_e.invoke({"solution_text": solution, "revenue_model": revenue_model, "trl": current_trl})

    # 10️⃣ Compile document
    doc = Document()
    for title, content in [
        ("The problem/market opportunity", analysis_a),
        ("The innovation: Solution/Product or Services (USP)", analysis_b),
        ("Market and Competition analysis", analysis_c),
        ("Broad impacts", analysis_d),
        ("Funding rationale and MVP", analysis_e)
    ]:
        doc.add_heading(title, level=1)
        doc.add_paragraph(content)
    output_file = "analyses.docx"
    doc.save(output_file)
    logger.info("Saved analyses document: %s", output_file)

    # 11️⃣ Send email with attachment
    try:
        with open(output_file, "rb") as f:
            encoded = base64.b64encode(f.read()).decode()
        message = Mail(
            from_email=SENDGRID_SENDER,
            to_emails=user_email,
            subject="Your AI-Generated Analyses",
            html_content="<p>Please find your requested analyses attached.</p>"
        )
        attachment = Attachment(
            FileContent(encoded),
            FileName(output_file),
            FileType("application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            Disposition("attachment")
        )
        message.attachment = attachment
        sg = SendGridAPIClient(SENDGRID_API_KEY)
        resp = sg.send(message)
        logger.info("Email sent; status code: %s", resp.status_code)
    except Exception as e:
        logger.error("Error sending email: %s", str(e))
        return JSONResponse(status_code=500, content={"detail": f"Analysis generated but email failed: {e}"})

    # 12️⃣ Return success with download link
    download_url = request.url._url.rstrip(request.url.path) + "/download"
    return {"status": "ok", "email_status": resp.status_code, "download_url": download_url}

@app.get("/download")
async def download_document():
    file_path = os.path.join(os.getcwd(), "analyses.docx")
    if not os.path.exists(file_path):
        return JSONResponse(status_code=404, content={"detail": "Document not found."})
    return FileResponse(path=file_path, filename="analyses.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=int(os.getenv("PORT", 8000)), log_level="info")
