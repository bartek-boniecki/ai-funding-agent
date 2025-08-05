# app.py
import os
import re
import uuid
import asyncio
from typing import Any, Dict, List, Optional, Tuple

from fastapi import FastAPI, Request, HTTPException
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from dotenv import load_dotenv
import logging
from datetime import datetime

# LLM + prompts (LangChain, modern style)
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser

# Web search providers
from langchain_community.utilities import SerpAPIWrapper

# Document creation
from docx import Document

# Email (SendGrid)
from sendgrid import SendGridAPIClient
from sendgrid.helpers.mail import Mail, Attachment, FileContent, FileName, FileType, Disposition
import base64

load_dotenv()

# ---------- Configuration & Logging ----------
REQUIRED_ENV = ["OPENAI_API_KEY", "SENDGRID_API_KEY", "SENDGRID_SENDER"]
missing = [k for k in REQUIRED_ENV if not os.getenv(k)]
if missing:
    raise RuntimeError(f"Missing env vars: {', '.join(missing)}")

SENDGRID_API_KEY = os.getenv("SENDGRID_API_KEY")
SENDGRID_SENDER = os.getenv("SENDGRID_SENDER")
SERPAPI_KEY = os.getenv("SERPAPI_API_KEY")  # optional; only needed if you want Google results

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
log = logging.getLogger("grant-agent")

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"]
)

# ---------- Models ----------
# Use a reliable, cost-efficient chat model. You can change to a reasoning model later if needed.
llm = ChatOpenAI(
    model=os.getenv("OPENAI_MODEL", "gpt-4o-mini"),
    temperature=0.1,  # slight creativity but still grounded
    timeout=120,
    max_retries=2,
)

parser = StrOutputParser()

# ---------- Prompts ----------
query_prompt = ChatPromptTemplate.from_messages([
    ("system",
     "You craft precise web search queries for professional market & problem research. "
     "Return ONLY the query, no extra text."),
    ("human",
     "Task: {task}\nContext: {context}\nReturn a single concise search query.")
])

# 4 research questions derived from your spec
RESEARCH_TASKS = {
    "problem_stats": "Find publicly available statistics that quantify the SCALE and URGENCY of the problem described.",
    "market": "Identify addressed market, market SIZE (global and/or regional as relevant), and KEY TRENDS for this product.",
    "novelty": "Identify companies with similar features and assess how NOVEL the proposed product is.",
    "revenues": "Find typical REVENUE STREAMS competitors generate from similar products."
}

analysis_problem_prompt = ChatPromptTemplate.from_messages([
    ("system",
     "You write grant-ready analyses (1,500–2,000 characters) with inline numeric evidence and cite sources as [1], [2], etc. "
     "Use only provided research notes and the user's description. Do not invent statistics."),
    ("human",
     "Title: The problem/market opportunity\n\n"
     "User problem description:\n{problem}\n\n"
     "Research notes (problem scale & urgency):\n{notes}\n\n"
     "Write a cohesive analysis between 1500 and 2000 characters. Include concrete figures and cite them inline as [n]. "
     "Do not truncate mid-sentence.")
])

analysis_innovation_prompt = ChatPromptTemplate.from_messages([
    ("system",
     "You write grant-ready analyses (1,500–2,000 characters) with inline citations [n]. Be precise and non-generic."),
    ("human",
     "Title: The innovation: Solution/Product or Services (USP)\n\n"
     "User solution:\n{solution}\n\n"
     "Unique features:\n{features}\n\n"
     "User-stated TRL: {trl}\n\n"
     "Research notes (competitors/alternatives & validation):\n{notes}\n\n"
     "Write 1,500–2,000 characters covering:\n"
     "1) Why this product is better than existing solutions (based on notes).\n"
     "2) EC TRL today and any validation/certification if indicated.\n"
     "3) Why now is the right time to launch.\n"
     "Cite sources [n]. Do not truncate mid-sentence.")
])

analysis_market_prompt = ChatPromptTemplate.from_messages([
    ("system",
     "You write grant-ready market analyses (1,500–2,000 chars) with inline citations [n]."),
    ("human",
     "Title: Market and Competition analysis\n\n"
     "User solution:\n{solution}\n\n"
     "Unique features:\n{features}\n\n"
     "User revenue ideas:\n{revenue}\n\n"
     "Research notes (market size & trends):\n{market_notes}\n\n"
     "Research notes (competitors & novelty):\n{novelty_notes}\n\n"
     "Research notes (competitor revenue streams):\n{rev_notes}\n\n"
     "Write 1,500–2,000 characters covering:\n"
     "(i) market size and key trends; (ii) potential to transform/create market;\n"
     "(iii) business model & revenue streams; (iv) why features will drive adoption;\n"
     "(v) advantages/disadvantages & success factors.\n"
     "Cite sources [n]. Do not truncate mid-sentence.")
])

analysis_impact_prompt = ChatPromptTemplate.from_messages([
    ("system",
     "You write grant-ready impact analyses (1,500–2,000 chars) with inline citations [n]."),
    ("human",
      "Title: Broad impacts\n\n"
      "User solution:\n{solution}\n\n"
      "Research notes (jobs, social/environmental/climate impacts):\n{notes}\n\n"
      "Write an analysis discussing potential societal/environmental/climate impacts and estimating jobs created. "
      "Length: 1,500–2,000 chars. Cite sources [n]. Do not truncate mid-sentence.")
])

analysis_funding_prompt = ChatPromptTemplate.from_messages([
    ("system",
     "You write candid funding analyses (1,500–2,000 chars) with practical reasoning."),
    ("human",
     "Title: Funding rationale and MVP\n\n"
     "User solution:\n{solution}\n\n"
     "User-stated TRL: {trl}\n\n"
     "User revenue ideas:\n{revenue}\n\n"
     "Write why raising funding is hard at this TRL per EC definitions, why an MVP is crucial now, and include any funding history if the user provided it in the form. "
     "Length: 1,500–2,000 chars. No citations required unless you choose to reference research already cited.")
])

# ---------- Helpers ----------
TRL_MAP = {
    "1": "TRL 1 – basic principles observed",
    "2": "TRL 2 – technology concept formulated",
    "3": "TRL 3 – experimental proof of concept",
    "4": "TRL 4 – technology validated in lab",
    "5": "TRL 5 – technology validated in relevant environment",
    "6": "TRL 6 – technology demonstrated in relevant environment",
    "7": "TRL 7 – system prototype demonstration in operational environment",
    "8": "TRL 8 – system complete and qualified",
    "9": "TRL 9 – actual system proven in operational environment"
}

def sanitize_trl(user_trl: str) -> str:
    digits = re.findall(r"\d", user_trl or "")
    if digits and digits[0] in TRL_MAP:
        return TRL_MAP[digits[0]]
    return user_trl.strip() or "Unspecified"

def ensure_length(name: str, text: str) -> None:
    n = len(text or "")
    if n < 100 or n > 2000:
        raise HTTPException(
            status_code=400,
            detail=f"Answer '{name}' must be between 100 and 2000 characters (got {n})."
        )

def extract_fillout_answers(payload: Dict[str, Any]) -> Dict[str, str]:
    """Map Fillout answers by question name, case-insensitive."""
    sub = payload.get("submission", {})
    questions = sub.get("questions", []) or []
    by_name = { (q.get("name") or "").strip().lower(): (q.get("value") or "").strip()
                for q in questions }

    # Expected names (set these EXACTLY in Fillout):
    expected = {
        "solution": "solution",
        "problem": "problem",
        "unique features": "features",
        "trl": "trl",
        "revenue streams": "revenue",
        "email": "email",
    }
    out = {}
    for key, alias in expected.items():
        val = by_name.get(key)
        if val is None:
            raise HTTPException(400, detail=f"Missing required field: '{key}' in Fillout submission.")
        out[alias] = val
    return out

# --- Web search (SerpAPI only here, but pluggable) ---
def init_serp() -> Optional[SerpAPIWrapper]:
    if not SERPAPI_KEY:
        return None
    return SerpAPIWrapper(serpapi_api_key=SERPAPI_KEY)

async def search_serp(serp: SerpAPIWrapper, query: str, k: int = 5) -> List[Dict[str, str]]:
    """Return a small, consistent list: [{'title','link','snippet'}]."""
    try:
        raw = serp.results(query)  # dict
        items = raw.get("organic_results", [])[:k]
        results = []
        for it in items:
            results.append({
                "title": it.get("title", ""),
                "link": it.get("link", ""),
                "snippet": it.get("snippet", "") or it.get("snippet_highlighted_words", [""])[0] if it else ""
            })
        return results
    except Exception as e:
        log.warning(f"SerpAPI error: {e}")
        return []

def format_notes_and_citations(results: List[Dict[str, str]]) -> Tuple[str, List[str]]:
    """
    Build readable notes with numbered citations and return (notes_text, citations_list).
    We only store distinct links to number them deterministically.
    """
    links: List[str] = []
    notes_lines: List[str] = []
    for it in results:
        url = it.get("link")
        if not url:
            continue
        if url not in links:
            links.append(url)
        idx = links.index(url) + 1
        title = it.get("title", "").strip()
        snip = it.get("snippet", "").strip()
        notes_lines.append(f"- {title}: {snip} [{idx}]")
    return ("\n".join(notes_lines).strip(), links)

def make_research_section_header(title: str) -> str:
    return f"\n---\n{title}\n"

# ---------- Routes ----------
@app.get("/health")
async def health():
    return {"status": "ok", "time": datetime.utcnow().isoformat() + "Z"}

@app.post("/webhook")
async def webhook(request: Request):
    payload = await request.json()
    try:
        answers = extract_fillout_answers(payload)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(400, detail=f"Malformed Fillout payload: {e}")

    # Validate lengths per spec
    ensure_length("solution", answers["solution"])
    ensure_length("problem", answers["problem"])
    ensure_length("unique features", answers["features"])
    ensure_length("trl", answers["trl"])
    ensure_length("revenue streams", answers["revenue"])

    email = answers["email"]
    if not re.match(r"[^@]+@[^@]+\.[^@]+", email):
        raise HTTPException(400, detail="Invalid email address.")

    # Generate refined queries (async)
    async def refine(task_key: str, context: str) -> str:
        prompt = query_prompt.format(task=RESEARCH_TASKS[task_key], context=context)
        return await (prompt | llm | parser).ainvoke({})

    # Prepare all 4 queries
    q_problem, q_market, q_novelty, q_revenue = await asyncio.gather(
        refine("problem_stats", answers["problem"]),
        refine("market", answers["solution"]),
        refine("novelty", answers["features"]),
        refine("revenues", answers["features"])
    )

    serp = init_serp()
    if not serp:
        log.warning("SERPAPI_API_KEY not set; research will be minimal.")
        research = {
            "problem": (f"(No web search: Missing SERPAPI_API_KEY)\nQuery: {q_problem}", []),
            "market": (f"(No web search: Missing SERPAPI_API_KEY)\nQuery: {q_market}", []),
            "novelty": (f"(No web search: Missing SERPAPI_API_KEY)\nQuery: {q_novelty}", []),
            "revenues": (f"(No web search: Missing SERPAPI_API_KEY)\nQuery: {q_revenue}", []),
        }
    else:
        # Do 4 searches in parallel (CPU-bound .results is light; keep it simple)
        loop = asyncio.get_event_loop()
        problem_res, market_res, novelty_res, revenue_res = await asyncio.gather(
            loop.run_in_executor(None, search_serp, serp, q_problem),
            loop.run_in_executor(None, search_serp, serp, q_market),
            loop.run_in_executor(None, search_serp, serp, q_novelty),
            loop.run_in_executor(None, search_serp, serp, q_revenue),
        )

        p_notes, p_links = format_notes_and_citations(problem_res)
        m_notes, m_links = format_notes_and_citations(market_res)
        n_notes, n_links = format_notes_and_citations(novelty_res)
        r_notes, r_links = format_notes_and_citations(revenue_res)

        research = {
            "problem": (p_notes, p_links),
            "market": (m_notes, m_links),
            "novelty": (n_notes, n_links),
            "revenues": (r_notes, r_links),
        }

    # Normalize TRL to EC description if possible
    trl_text = sanitize_trl(answers["trl"])

    # Generate analyses (async)
    a_problem = (analysis_problem_prompt | llm | parser).ainvoke({
        "problem": answers["problem"],
        "notes": research["problem"][0]
    })
    a_innovation = (analysis_innovation_prompt | llm | parser).ainvoke({
        "solution": answers["solution"],
        "features": answers["features"],
        "trl": trl_text,
        "notes": research["novelty"][0] or research["market"][0]
    })
    a_market = (analysis_market_prompt | llm | parser).ainvoke({
        "solution": answers["solution"],
        "features": answers["features"],
        "revenue": answers["revenue"],
        "market_notes": research["market"][0],
        "novelty_notes": research["novelty"][0],
        "rev_notes": research["revenues"][0]
    })
    a_impact = (analysis_impact_prompt | llm | parser).ainvoke({
        "solution": answers["solution"],
        "notes": research["revenues"][0] or research["market"][0]
    })
    a_funding = (analysis_funding_prompt | llm | parser).ainvoke({
        "solution": answers["solution"],
        "trl": trl_text,
        "revenue": answers["revenue"],
    })

    a_txt, b_txt, c_txt, d_txt, e_txt = await asyncio.gather(
        a_problem, a_innovation, a_market, a_impact, a_funding
    )

    # Build references once across all sections (dedupe preserving order)
    all_links: List[str] = []
    for sec in ["problem", "market", "novelty", "revenues"]:
        for url in research[sec][1]:
            if url not in all_links:
                all_links.append(url)

    # Create the Word document with per-section headings and a References section
    file_id = uuid.uuid4().hex[:8]
    filename = f"analyses_{file_id}.docx"
    doc = Document()

    doc.add_heading("The problem/market opportunity", level=1)
    doc.add_paragraph(a_txt)

    doc.add_heading("The innovation: Solution/Product or Services (USP)", level=1)
    doc.add_paragraph(b_txt)

    doc.add_heading("Market and Competition analysis", level=1)
    doc.add_paragraph(c_txt)

    doc.add_heading("Broad impacts", level=1)
    doc.add_paragraph(d_txt)

    doc.add_heading("Funding rationale and MVP", level=1)
    doc.add_paragraph(e_txt)

    # Optional: add raw research notes (helpful for transparency)
    doc.add_heading("Appendix: Research notes", level=1)
    for title, key in [
        ("Problem scale & urgency", "problem"),
        ("Market size & trends", "market"),
        ("Competitors & novelty", "novelty"),
        ("Competitor revenue streams", "revenues")
    ]:
        doc.add_heading(title, level=2)
        notes_text = research[key][0] or "(No results)"
        for line in (notes_text.split("\n") if notes_text else []):
            doc.add_paragraph(line)

    if all_links:
        doc.add_heading("References", level=1)
        for i, url in enumerate(all_links, 1):
            doc.add_paragraph(f"[{i}] {url}")

    doc.save(filename)
    log.info(f"Saved document: {filename}")

    # Email the document
    try:
        with open(filename, "rb") as f:
            data_b64 = base64.b64encode(f.read()).decode()
        message = Mail(
            from_email=SENDGRID_SENDER,
            to_emails=email,
            subject="Your grant analyses (AI-generated with sources)",
            html_content="<p>Your analyses are attached. We also included an Appendix and References.</p>"
        )
        attachment = Attachment(
            FileContent(data_b64),
            FileName(os.path.basename(filename)),
            FileType("application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            Disposition("attachment")
        )
        message.attachment = attachment
        resp = SendGridAPIClient(SENDGRID_API_KEY).send(message)
        email_status = resp.status_code
        log.info(f"Email sent to {email}; status {email_status}")
    except Exception as e:
        log.error(f"SendGrid error: {e}")
        email_status = None

    # Return a JSON with a unique download link
    base = str(request.url).replace("/webhook", "")
    download_url = f"{base}/download/{filename}"
    return JSONResponse({"status": "ok", "email_status": email_status, "download_url": download_url})

@app.get("/download/{fname}")
async def download(fname: str):
    path = os.path.join(os.getcwd(), fname)
    if not os.path.exists(path):
        raise HTTPException(404, detail="Document not found")
    return FileResponse(
        path=path,
        filename=fname,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=int(os.getenv("PORT", "8000")), log_level="info")
